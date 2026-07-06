import os
import shutil
from fastapi import APIRouter, HTTPException, Header, UploadFile, File
from typing import Optional
from pydantic import BaseModel
from app.models.prompt import get_all_prompts, create_prompt, update_prompt, delete_prompt
from app.database import get_db_connection
import psycopg2.extras

ASSETS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "assets")
os.makedirs(ASSETS_DIR, exist_ok=True)

router = APIRouter()

class PromptBase(BaseModel):
    name: str
    prompt_type: str
    content: str
    description: Optional[str] = None
    is_active: Optional[bool] = True

class PromptCreate(BaseModel):
    name: str
    content: str
    description: Optional[str] = None
    followup_1: Optional[str] = None
    followup_2: Optional[str] = None
    followup_3: Optional[str] = None
    subject: Optional[str] = None
    cc: Optional[str] = None
    followup_count: Optional[int] = 3

class PromptUpdate(BaseModel):
    name: Optional[str] = None
    prompt_type: Optional[str] = None
    content: Optional[str] = None
    description: Optional[str] = None
    is_active: Optional[bool] = None
    followup_1: Optional[str] = None
    followup_2: Optional[str] = None
    followup_3: Optional[str] = None
    subject: Optional[str] = None
    cc: Optional[str] = None
    followup_count: Optional[int] = None

@router.get("/prompts")
def list_prompts():
    return get_all_prompts()

@router.post("/prompts")
def add_prompt(prompt: PromptBase):
    prompt_id = create_prompt(
        prompt.name, 
        prompt.prompt_type, 
        prompt.content, 
        prompt.description, 
        prompt.is_active
    )
    return {"id": prompt_id, "message": "Prompt created successfully"}

@router.post("/custom-draft-templates")
def create_custom_template(tpl: PromptCreate, user_id: Optional[str] = Header(None, alias="X-User-Id")):
    """Create a custom draft template with follow-ups, owned by the current user."""
    owner_username = None
    if user_id:
        try:
            conn = get_db_connection()
            cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
            cur.execute("SELECT username, full_name FROM users WHERE id = %s", (int(user_id),))
            user_row = cur.fetchone()
            cur.close()
            conn.close()
            if user_row:
                uname = str(user_row['username'] or '').lower()
                fname = str(user_row['full_name'] or '').lower()
                owner_username = uname.split('.')[0] or fname.split()[0] if fname else uname
        except Exception:
            pass
    prompt_id = create_prompt(
        name=tpl.name,
        prompt_type='CUSTOM_DRAFT',
        content=tpl.content,
        description=tpl.description,
        is_active=True,
        owner_username=owner_username,
        followup_1=tpl.followup_1,
        followup_2=tpl.followup_2,
        followup_3=tpl.followup_3,
        subject=tpl.subject,
        cc=tpl.cc,
        followup_count=tpl.followup_count
    )
    return {"id": prompt_id, "message": "Custom template created successfully"}

@router.put("/prompts/{prompt_id}")
def edit_prompt(prompt_id: int, prompt_data: PromptUpdate):
    success = update_prompt(prompt_id, prompt_data.dict(exclude_unset=True))
    if not success:
        raise HTTPException(status_code=404, detail="Prompt not found")
    return {"message": "Prompt updated successfully"}

ALLOWED_IMAGE_TYPES = {
    'image/png': '.png',
    'image/jpeg': '.jpg',
    'image/jpg': '.jpg',
    'image/gif': '.gif',
    'image/webp': '.webp',
    'image/svg+xml': '.svg',
}

ALLOWED_DOCUMENT_TYPES = {
    'application/pdf': '.pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
    'application/msword': '.doc',
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx',
    'application/vnd.ms-excel': '.xls',
}

@router.post("/upload-image")
def upload_image(file: UploadFile = File(...)):
    """Upload an image to the assets folder and return its URL path."""
    if file.content_type not in ALLOWED_IMAGE_TYPES:
        raise HTTPException(status_code=400, detail=f"Invalid image type: {file.content_type}. Allowed: PNG, JPG, GIF, WebP, SVG")
    ext = ALLOWED_IMAGE_TYPES[file.content_type]
    # Generate unique filename
    import time, random
    ts = int(time.time() * 1000)
    rn = random.randint(1000, 9999)
    dest_name = f"upload_{ts}_{rn}{ext}"
    dest_path = os.path.join(ASSETS_DIR, dest_name)
    with open(dest_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    backend_url = os.getenv("BACKEND_URL", "http://127.0.0.1:8000").rstrip("/")
    return {"filename": dest_name, "url": f"{backend_url}/assets/{dest_name}"}

@router.post("/upload-file")
def upload_file(file: UploadFile = File(...)):
    """Upload a document file (PDF, DOCX, XLSX) to the assets folder and return its URL."""
    if file.content_type not in ALLOWED_DOCUMENT_TYPES:
        raise HTTPException(status_code=400, detail=f"Invalid document type: {file.content_type}. Allowed: PDF, DOC, DOCX, XLS, XLSX")
    ext = ALLOWED_DOCUMENT_TYPES[file.content_type]
    import time, random
    ts = int(time.time() * 1000)
    rn = random.randint(1000, 9999)
    safe_name = "".join(c for c in file.filename.rsplit('.', 1)[0] if c.isalnum() or c in ' _-')[:60]
    dest_name = f"{safe_name}_{ts}_{rn}{ext}"
    dest_path = os.path.join(ASSETS_DIR, dest_name)
    with open(dest_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    backend_url = os.getenv("BACKEND_URL", "http://127.0.0.1:8000").rstrip("/")
    return {"filename": dest_name, "url": f"{backend_url}/assets/{dest_name}"}

@router.post("/upload-signature-doc")
def upload_signature_doc(file: UploadFile = File(...)):
    """Upload a Word/PDF document and extract formatted text to use as signature.
    
    Preserves: bold, italic, headings, hyperlinks, images, lists.
    Images from DOCX are extracted to the assets folder and embedded as data URIs
    so they render reliably in signature preview and sent emails.
    """
    import tempfile, re, uuid
    name_lower = (file.filename or '').lower()
    if not any(name_lower.endswith(ext) for ext in ['.docx', '.pdf', '.doc']):
        raise HTTPException(status_code=400, detail="Only DOCX, PDF, and DOC files are allowed")
    ext = name_lower.rsplit('.', 1)[-1]
    tmp = tempfile.NamedTemporaryFile(suffix=f'.{ext}', delete=False)
    try:
        shutil.copyfileobj(file.file, tmp)
        tmp.close()

        if ext == 'docx':
            from docx import Document
            from docx.oxml.ns import qn
            from docx.oxml import parse_xml
            doc = Document(tmp.name)
            lines = []
            img_map = {}
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.reltype:
                    image_blob = rel.target_part.blob
                    ext_map = {'image/png': '.png', 'image/jpeg': '.jpg', 'image/jpg': '.jpg', 'image/gif': '.gif', 'image/bmp': '.bmp', 'image/tiff': '.tiff', 'image/webp': '.webp'}
                    mime = rel.target_part.content_type
                    img_ext = ext_map.get(mime, '.png')
                    img_name = f"sig_{uuid.uuid4().hex[:12]}{img_ext}"
                    img_path = os.path.join(ASSETS_DIR, img_name)
                    with open(img_path, "wb") as imgf:
                        imgf.write(image_blob)
                    import base64 as _b64
                    b64data = _b64.b64encode(image_blob).decode()
                    data_uri = f"data:{mime};base64,{b64data}"
                    img_map[rel_id] = f"![image]({data_uri})"

            for table in doc.tables:
                table_html = '<table style="border-collapse:collapse;width:100%;border:1px solid #d0d0d0;">'
                for row in table.rows:
                    table_html += '<tr>'
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        table_html += f'<td style="border:1px solid #d0d0d0;padding:4px 8px;text-align:left;vertical-align:top;">{cell_text}</td>'
                    table_html += '</tr>'
                table_html += '</table>'
                lines.append(table_html)
                lines.append('')

            for para in doc.paragraphs:
                txt = para.text.strip()
                style = para.style.name.lower() if para.style else ''
                alignment = para.paragraph_format.alignment

                for run in para.runs:
                    for child in run._element:
                        if child.tag.endswith('}drawing'):
                            blip_fill = child.find('.//' + qn('a:blip'))
                            if blip_fill is not None:
                                embed_id = blip_fill.get(qn('r:embed'))
                                if embed_id and embed_id in img_map:
                                    lines.append('')
                                    lines.append(img_map[embed_id])
                                    lines.append('')

                if not txt:
                    if para.paragraph_format.space_before or para.paragraph_format.space_after:
                        lines.append('<br>')
                    continue

                para_text = _build_run_text(para)

                if alignment is not None:
                    align_map = {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}
                    align_val = align_map.get(alignment, 'left')
                    if align_val != 'left':
                        para_text = f'<div style="text-align:{align_val};">{para_text}</div>'

                if 'heading' in style:
                    level = re.search(r'heading\s*(\d+)', style)
                    if level:
                        prefix = '#' * int(level.group(1))
                        lines.append(f'{prefix} {para_text}')
                    else:
                        lines.append(f'**{para_text}**')
                elif para.style and 'list' in style:
                    lines.append(f'- {para_text}')
                else:
                    lines.append(para_text)

            text = '\n'.join(lines).strip()
        elif ext == 'pdf':
            from PyPDF2 import PdfReader
            reader = PdfReader(tmp.name)
            lines = []
            for page in reader.pages:
                txt = page.extract_text().strip()
                if txt:
                    lines.append(txt)
            text = '\n'.join(lines)
        else:
            text = f'[Uploaded file: {file.filename}]({os.getenv("BACKEND_URL", "http://127.0.0.1:8000").rstrip("/")}/assets/{uuid.uuid4().hex}.pdf)'
        return {"text": text, "filename": file.filename}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract text: {str(e)}")
    finally:
        os.unlink(tmp.name)


def _build_run_text(para) -> str:
    """Build a markdown string from paragraph runs, preserving per-run formatting and hyperlinks.
    
    Supports: bold, italic, underline, font name, font size, font color.
    Adjacent runs inside the same hyperlink are merged into a single markdown link.
    """
    parts = []
    link_map = {}
    for hyperlink in para.hyperlinks:
        url = getattr(hyperlink, 'url', None) or getattr(hyperlink, 'address', None) or ''
        link_map[id(hyperlink._hyperlink)] = url

    link_buf = []
    link_url = None

    def _flush_link():
        nonlocal link_buf, link_url
        if link_buf:
            txt = ''.join(link_buf)
            parts.append(f'[{txt}]({link_url})')
            link_buf = []
            link_url = None

    for run in para.runs:
        txt = run.text
        if not txt:
            _flush_link()
            continue

        parent = run._element.getparent()
        run_link_url = None
        if parent is not None and id(parent) in link_map:
            run_link_url = link_map[id(parent)]

        if link_url is not None and (run_link_url is None or run_link_url != link_url):
            _flush_link()
        if link_url is None and run_link_url is not None:
            link_url = run_link_url

        # Build rich formatting: extract all font properties
        styles = []
        has_extra = False

        # Underline
        if run.font.underline:
            has_extra = True
            styles.append('text-decoration:underline')

        # Font size (in points)
        if run.font.size:
            sz = run.font.size.pt
            if sz:
                has_extra = True
                styles.append(f'font-size:{int(sz)}pt')

        # Font color
        if run.font.color:
            try:
                rgb = run.font.color.rgb
                if rgb:
                    has_extra = True
                    styles.append(f'color:#{rgb}')
            except Exception:
                pass

        # Font name
        if run.font.name:
            has_extra = True
            fn = run.font.name
            if ' ' in fn and not fn.startswith('"'):
                fn = f'"{fn}"'
            styles.append(f'font-family:{fn}')

        # Apply markdown formatting
        if run.bold:
            txt = f'**{txt}**'
        if run.italic:
            txt = f'*{txt}*'

        # If there are extra font properties, wrap in a <span>
        if has_extra:
            txt = f'<span style="{"; ".join(styles)};">{txt}</span>'

        if link_url is not None:
            link_buf.append(txt)
        else:
            parts.append(txt)

    _flush_link()
    return ''.join(parts)

@router.post("/prompts/{prompt_id}/attachment")
def upload_prompt_attachment(prompt_id: int, file: UploadFile = File(...)):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files allowed")
    ext = os.path.splitext(file.filename)[1]
    dest_name = f"prompt_{prompt_id}_attachment{ext}"
    dest_path = os.path.join(ASSETS_DIR, dest_name)
    with open(dest_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("UPDATE prompts SET attachment_file = %s WHERE id = %s", (dest_name, prompt_id))
    conn.commit()
    cur.close()
    conn.close()
    return {"filename": dest_name, "message": "Attachment uploaded"}

@router.delete("/prompts/{prompt_id}")
def remove_prompt(prompt_id: int, user_id: Optional[str] = Header(None, alias="X-User-Id")):
    """Delete a prompt — only if the user owns it."""
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
    # Fetch the prompt to verify ownership
    cur.execute("SELECT owner_username FROM prompts WHERE id = %s", (prompt_id,))
    prompt = cur.fetchone()
    if not prompt:
        cur.close()
        conn.close()
        raise HTTPException(status_code=404, detail="Prompt not found")
    owner = prompt['owner_username']
    if owner and user_id:
        try:
            cur.execute("SELECT username, full_name FROM users WHERE id = %s", (int(user_id),))
            user_row = cur.fetchone()
            if user_row:
                uname = str(user_row['username'] or '').lower()
                fname = str(user_row['full_name'] or '').lower()
                current_user = uname.split('.')[0] or fname.split()[0] if fname else uname
                if current_user != owner:
                    cur.close()
                    conn.close()
                    raise HTTPException(status_code=403, detail="You can only delete your own templates")
        except HTTPException:
            raise
        except Exception:
            pass
    elif owner and not user_id:
        cur.close()
        conn.close()
        raise HTTPException(status_code=403, detail="Authentication required to delete this template")
    cur.close()
    conn.close()
    success = delete_prompt(prompt_id)
    if not success:
        raise HTTPException(status_code=404, detail="Prompt not found")
    return {"message": "Template deleted successfully"}
